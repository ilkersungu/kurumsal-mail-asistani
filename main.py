import streamlit as st
from huggingface_hub import InferenceClient
from io import BytesIO
from docx import Document

# --- 1. AYARLAR ---
# Şifreyi Streamlit'in gizli kasasından (Secrets) çekiyoruz
try:
    HF_TOKEN = st.secrets["HF_TOKEN"]
except:
    # Eğer bilgisayarında çalıştırıyorsan burayı açıp kendi şifreni yazabilirsin test için
    # Ama GitHub'a atarken burayı gizli tutmak en iyisidir.
    st.error("⚠️ API Key bulunamadı! Streamlit Secrets ayarlarını kontrol et.")
    st.stop()

repo_id = "Qwen/Qwen2.5-7B-Instruct"
client = InferenceClient(model=repo_id, token=HF_TOKEN)

# --- 2. SAYFA YAPISI ---
st.set_page_config(page_title="Kurumsal Asistan Pro", page_icon="🏢", layout="wide")

# --- YARDIMCI FONKSİYON: WORD DOSYASI OLUŞTURUCU ---
def word_dosyasi_olustur(mail_metni, baslik, gonderen, sirket):
    doc = Document()
    # Başlık Ekle
    doc.add_heading(baslik, 0)
    # Bilgi satırı
    doc.add_paragraph(f"Oluşturan: {gonderen} | {sirket}")
    doc.add_paragraph("-" * 50)
    # Ana Metin
    doc.add_paragraph(mail_metni)
    # Dosyayı hafızaya kaydet (Diske değil)
    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio

# --- 3. VERİTABANI ---
sirket_senaryolari = {
    "👥 İnsan Kaynakları (İK)": ["Yıllık izin talebi", "Maaş zammı talebi", "İstifa dilekçesi", "Personel alımı"],
    "💻 Bilgi İşlem (IT)": ["Bilgisayar yavaş", "İnternet yok", "VPN yetkisi", "Yeni Ekipman Talebi"],
    "💰 Muhasebe & Finans": ["Maaş yatmadı", "Avans talebi", "Fatura teslimi"],
    "🤝 Satış & Pazarlama": ["Toplantı özeti", "Teklif maili", "Müşteri şikayeti yanıtlama"],
    "⚖️ Hukuk & Sözleşmeler": ["Sözleşme taslağı", "İhtarname taslağı"],
    "🏢 İdari İşler": ["Ofis temizliği", "Klima sorunu", "Duyuru"]
}

# --- 4. SOL MENÜ ---
with st.sidebar:
    st.header("👤 Ayarlar")
    gonderen_ad = st.text_input("Adınız Soyadınız", placeholder="Örn: İlker Yılmaz")
    unvan = st.text_input("Unvanınız", placeholder="Örn: Uzman")
    sirket_adi = st.text_input("Şirket Adı", placeholder="Örn: Yılmaz A.Ş.")
    
    st.markdown("---")
    ton = st.select_slider("Üslup Seçin:", options=["Çok Resmi", "Kurumsal", "Nazik", "Sert/Net", "Arkadaşça"], value="Kurumsal")

# --- 5. ANA EKRAN ---
st.title("🏢 Kurumsal İletişim Asistanı")
st.markdown("Word çıktısı alabileceğiniz profesyonel mail oluşturucu.")

col1, col2 = st.columns(2)
with col1:
    secilen_departman = st.selectbox("Departman:", list(sirket_senaryolari.keys()))
    kime = st.text_input("Alıcı Adı:", placeholder="Örn: Mehmet Bey")
with col2:
    secilen_konu = st.selectbox("Konu:", sirket_senaryolari[secilen_departman])
    tarih_ekle = st.checkbox("Tarih Ekle?")
    tarih_str = f"Tarih: {st.date_input('Tarih:')}" if tarih_ekle else ""

detaylar = st.text_area("✍️ Ekstra Detaylar:", height=100)

# --- BUTON VE İŞLEM ---
if st.button("🚀 Maili Oluştur", use_container_width=True):
    if not gonderen_ad or not sirket_adi:
        st.error("Lütfen sol menüden bilgilerinizi girin!")
    elif not kime:
        st.warning("Alıcı ismini girmelisiniz.")
    else:
        # Prompt Hazırlığı
        messages = [
            {"role": "system", "content": f"Sen {sirket_adi} şirketinde {unvan} olan {gonderen_ad}. '{ton}' tonunda Türkçe mail yaz. Asla yer tutucu (köşeli parantez) bırakma. Sadece mail içeriğini ver."},
            {"role": "user", "content": f"Alıcı: {kime}\nKonu: {secilen_konu}\nDetay: {detaylar} {tarih_str}"}
        ]
        
        with st.spinner('Yapay zeka yazıyor...'):
            try:
                response = client.chat_completion(messages, max_tokens=800)
                mail_metni = response.choices[0].message.content
                
                # Ekrana Yazdır
                st.success("✅ Mail Hazır!")
                st.text_area("Önizleme:", value=mail_metni, height=400)
                
                # --- YENİ ÖZELLİK: WORD İNDİRME ---
                st.markdown("### 📥 İndirme Seçenekleri")
                
                # Word dosyasını oluşturuyoruz
                word_data = word_dosyasi_olustur(mail_metni, secilen_konu, gonderen_ad, sirket_adi)
                
                # İndirme Butonu
                st.download_button(
                    label="📄 Word Dosyası Olarak İndir (.docx)",
                    data=word_data,
                    file_name=f"{secilen_konu}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
                
            except Exception as e:
                st.error(f"Hata: {e}")